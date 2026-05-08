# app/blueprints/ocorrencia/routes.py

import logging
import re
from datetime import datetime
from functools import wraps

import pytz
from flask import (Blueprint, current_app, flash, jsonify, redirect,
                   render_template, request, url_for, send_file)
import io
from flask_login import current_user, login_required
from flask_cors import cross_origin
from flask_wtf import csrf

from app import db
from app.forms import OcorrenciaForm
from app.models import (Colaborador, Condominio, Ocorrencia, OcorrenciaTipo,
                        OrgaoPublico, User)
from app.services import ocorrencia_service
from app.services.email_patrimonial_format_service import EmailPatrimonialFormatService
from app.utils.classificador import classificar_ocorrencia

logger = logging.getLogger(__name__)

ocorrencia_bp = Blueprint(
    "ocorrencia", __name__, template_folder="templates", url_prefix="/ocorrencias"
)


# --- Funções Auxiliares ---
def local_to_utc(dt_naive, timezone_str=None):
    """Converte datetime ingênuo (sem tzinfo) local para datetime UTC com tzinfo."""
    ## [MELHORIA] Usa o timezone da configuração da app, com um fallback.
    if not timezone_str:
        timezone_str = current_app.config.get("DEFAULT_TIMEZONE", "America/Sao_Paulo")

    local_tz = pytz.timezone(timezone_str)
    aware_local = local_tz.localize(dt_naive)
    utc_dt = aware_local.astimezone(pytz.utc)
    return utc_dt


def populate_ocorrencia_form_choices(form):
    form.condominio_id.choices = [("", "-- Selecione um Condomínio --")] + [
        (str(c.id), c.nome) for c in Condominio.query.order_by("nome").all()
    ]
    form.ocorrencia_tipo_id.choices = [("", "-- Selecione um Tipo --")] + [
        (str(t.id), t.nome) for t in OcorrenciaTipo.query.order_by("nome").all()
    ]
    form.orgaos_acionados.choices = [
        (str(o.id), o.nome) for o in OrgaoPublico.query.order_by("nome").all()
    ]
    form.colaboradores_envolvidos.choices = [
        (str(col.id), col.nome_completo)
        for col in Colaborador.query.filter_by(status="Ativo")
        .order_by("nome_completo")
        .all()
    ]
    form.supervisor_id.choices = [("", "-- Selecione um Supervisor --")] + [
        (str(s.id), s.username)
        for s in User.query.filter_by(is_supervisor=True, is_approved=True)
        .order_by("username")
        .all()
    ]


def pode_editar_ocorrencia(f):
    @wraps(f)
    def decorated_function(ocorrencia_id, *args, **kwargs):
        ocorrencia = db.get_or_404(Ocorrencia, ocorrencia_id)
        if not (
            current_user.is_admin
            or current_user.id == ocorrencia.registrado_por_user_id
        ):
            flash("Você não tem permissão para editar esta ocorrência.", "danger")
            return redirect(
                url_for("ocorrencia.detalhes_ocorrencia", ocorrencia_id=ocorrencia.id)
            )
        return f(ocorrencia_id, *args, **kwargs)

    return decorated_function


# --- Rotas ---
@ocorrencia_bp.route("/historico")
@login_required
def listar_ocorrencias():
    page = request.args.get("page", 1, type=int)

    # Coleta todos os argumentos de filtro em um dicionário
    filters = {
        "status": request.args.get("status", ""),
        "condominio_id": request.args.get("condominio_id", type=int),
        "supervisor_id": request.args.get("supervisor_id", type=int),
        "tipo_id": request.args.get("tipo_id", type=int),
        "data_inicio": request.args.get("data_inicio", ""),
        "data_fim": request.args.get("data_fim", ""),
        "texto_relatorio": request.args.get("texto_relatorio", ""),
    }

    from app.models.vw_ocorrencias_detalhadas import VWOcorrenciasDetalhadas
    query = VWOcorrenciasDetalhadas.query

    ## [REMOVIDO] Bloco inteiro de IFs para filtragem manual foi removido daqui.
    # if selected_status: query = query.filter(Ocorrencia.status == selected_status)
    # ... e todos os outros ifs ...

    ## [ADICIONADO] Chamada única para a função de serviço centralizada.
    query = ocorrencia_service.apply_ocorrencia_filters(query, filters)

    from app.models.vw_ocorrencias_detalhadas import VWOcorrenciasDetalhadas
    ocorrencias_pagination = query.order_by(
        VWOcorrenciasDetalhadas.data_hora_ocorrencia.desc()
    ).paginate(page=page, per_page=15, error_out=False)

    # Carrega dados para preencher os formulários de filtro no template
    condominios = Condominio.query.order_by(Condominio.nome).all()
    supervisors = (
        User.query.filter_by(is_supervisor=True, is_approved=True)
        .order_by(User.username)
        .all()
    )
    tipos_ocorrencia = OcorrenciaTipo.query.order_by(OcorrenciaTipo.nome).all()
    status_list = [
        "Registrada",
        "Em Andamento",
        "Concluída",
        "Pendente",
        "Rejeitada"
    ]

    filter_args = {k: v for k, v in request.args.items() if k != "page"}

    return render_template(
        "ocorrencia/list.html",
        title="Histórico de Ocorrências",
        ocorrencias_pagination=ocorrencias_pagination,
        condominios=condominios,
        supervisors=supervisors,
        tipos_ocorrencia=tipos_ocorrencia,
        status_list=status_list,
        selected_status=filters["status"],
        selected_condominio_id=filters["condominio_id"],
        selected_supervisor_id=filters["supervisor_id"],
        selected_tipo_id=filters["tipo_id"],
        selected_data_inicio=filters["data_inicio"],
        selected_data_fim=filters["data_fim"],
        filter_args=filter_args,
        texto_relatorio=filters["texto_relatorio"],
    )


## O restante do arquivo (registrar, editar, deletar, analisar) permanece o mesmo.
@ocorrencia_bp.route("/registrar", methods=["GET", "POST"])
@login_required
def registrar_ocorrencia():
    form = OcorrenciaForm()
    populate_ocorrencia_form_choices(form)

    if request.method == "GET":
        if request.args.get("relatorio_final"):
            form.relatorio_final.data = request.args.get("relatorio_final")
        form.data_hora_ocorrencia.data = datetime.now()

    if form.validate_on_submit():
        try:
            tipo_ocorrencia_id = form.ocorrencia_tipo_id.data
            if form.novo_tipo_ocorrencia.data:
                tipo_existente = OcorrenciaTipo.query.filter(
                    OcorrenciaTipo.nome.ilike(form.novo_tipo_ocorrencia.data.strip())
                ).first()
                if tipo_existente:
                    tipo_ocorrencia_id = tipo_existente.id
                else:
                    novo_tipo = OcorrenciaTipo(
                        nome=form.novo_tipo_ocorrencia.data.strip()
                    )
                    db.session.add(novo_tipo)
                    db.session.flush()
                    tipo_ocorrencia_id = novo_tipo.id
            
            # Validação adicional para garantir que tipo_ocorrencia_id não seja None
            if not tipo_ocorrencia_id:
                # Busca um tipo padrão se nenhum foi selecionado
                tipo_padrao = OcorrenciaTipo.query.filter_by(nome="verificação").first()
                if not tipo_padrao:
                    # Cria o tipo padrão se não existir
                    tipo_padrao = OcorrenciaTipo(nome="verificação")
                    db.session.add(tipo_padrao)
                    db.session.flush()
                tipo_ocorrencia_id = tipo_padrao.id

            utc_datetime = None
            if form.data_hora_ocorrencia.data:
                utc_datetime = local_to_utc(form.data_hora_ocorrencia.data)

            nova_ocorrencia = Ocorrencia(
                condominio_id=form.condominio_id.data,
                data_hora_ocorrencia=utc_datetime,
                turno=form.turno.data,
                relatorio_final=form.relatorio_final.data,
                status=form.status.data,
                endereco_especifico=form.endereco_especifico.data,
                ocorrencia_tipo_id=tipo_ocorrencia_id,
                registrado_por_user_id=current_user.id,
                supervisor_id=form.supervisor_id.data,
            )
            if form.orgaos_acionados.data:
                orgaos = OrgaoPublico.query.filter(
                    OrgaoPublico.id.in_(form.orgaos_acionados.data)
                ).all()
                nova_ocorrencia.orgaos_acionados.extend(orgaos)
            if form.colaboradores_envolvidos.data:
                colaboradores = Colaborador.query.filter(
                    Colaborador.id.in_(form.colaboradores_envolvidos.data)
                ).all()
                nova_ocorrencia.colaboradores_envolvidos.extend(colaboradores)

            db.session.add(nova_ocorrencia)
            db.session.commit()
            flash("Ocorrência registrada com sucesso!", "success")
            return redirect(
                url_for(
                    "ocorrencia.detalhes_ocorrencia", ocorrencia_id=nova_ocorrencia.id
                )
            )
        except Exception as e:
            db.session.rollback()
            logger.error(f"Erro ao salvar nova ocorrência: {e}", exc_info=True)
            flash(f"Erro ao salvar a ocorrência: {e}", "danger")

    return render_template(
        "ocorrencia/form_direto.html", title="Registrar Nova Ocorrência", form=form
    )


@ocorrencia_bp.route("/detalhes/<int:ocorrencia_id>")
@login_required
def detalhes_ocorrencia(ocorrencia_id):
    ocorrencia = db.get_or_404(Ocorrencia, ocorrencia_id)
    return render_template(
        "ocorrencia/details.html",
        title=f"Detalhes da Ocorrência #{ocorrencia.id}",
        ocorrencia=ocorrencia,
    )


@ocorrencia_bp.route("/editar/<int:ocorrencia_id>", methods=["GET", "POST"])
@login_required
@pode_editar_ocorrencia
def editar_ocorrencia(ocorrencia_id):
    ocorrencia = db.get_or_404(Ocorrencia, ocorrencia_id)
    form = OcorrenciaForm(obj=ocorrencia)
    populate_ocorrencia_form_choices(form)

    if request.method == "GET":
        form.colaboradores_envolvidos.data = [
            col.id for col in ocorrencia.colaboradores_envolvidos
        ]
        form.orgaos_acionados.data = [org.id for org in ocorrencia.orgaos_acionados]

    if form.validate_on_submit():
        try:
            # Popula apenas os campos simples, exceto os relacionamentos
            campos_simples = [
                "condominio_id", "data_hora_ocorrencia", "turno", "relatorio_final",
                "status", "endereco_especifico", "ocorrencia_tipo_id", "supervisor_id"
            ]
            for campo in campos_simples:
                if hasattr(form, campo) and hasattr(ocorrencia, campo):
                    setattr(ocorrencia, campo, getattr(form, campo).data)

            # Lida com o datetime
            if form.data_hora_ocorrencia.data:
                ocorrencia.data_hora_ocorrencia = local_to_utc(
                    form.data_hora_ocorrencia.data
                )

            # Lida com relacionamentos Many-to-Many
            colaboradores = Colaborador.query.filter(
                Colaborador.id.in_(form.colaboradores_envolvidos.data or [])
            ).all()
            ocorrencia.colaboradores_envolvidos = colaboradores

            orgaos = OrgaoPublico.query.filter(
                OrgaoPublico.id.in_(form.orgaos_acionados.data or [])
            ).all()
            ocorrencia.orgaos_acionados = orgaos

            db.session.commit()
            flash("Ocorrência atualizada com sucesso!", "success")
            return redirect(
                url_for("ocorrencia.detalhes_ocorrencia", ocorrencia_id=ocorrencia.id)
            )
        except Exception as e:
            db.session.rollback()
            logger.error(f"Erro ao atualizar ocorrência: {e}", exc_info=True)
            flash(f"Erro ao atualizar a ocorrência: {e}", "danger")

    return render_template(
        "ocorrencia/form_direto.html",
        title=f"Editar Ocorrência #{ocorrencia.id}",
        form=form,
    )


@ocorrencia_bp.route("/deletar/<int:ocorrencia_id>", methods=["POST"])
@login_required
def deletar_ocorrencia(ocorrencia_id):
    ocorrencia = db.get_or_404(Ocorrencia, ocorrencia_id)
    if not (
        current_user.is_admin or current_user.id == ocorrencia.registrado_por_user_id
    ):
        flash("Você não tem permissão para deletar esta ocorrência.", "danger")
        return redirect(url_for("ocorrencia.listar_ocorrencias"))
    try:
        db.session.delete(ocorrencia)
        db.session.commit()
        flash("Ocorrência deletada com sucesso.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Erro ao deletar a ocorrência: {e}", "danger")
    return redirect(url_for("ocorrencia.listar_ocorrencias"))


@ocorrencia_bp.route("/aprovar/<int:ocorrencia_id>", methods=["POST"])
@login_required
def aprovar_ocorrencia(ocorrencia_id):
    """Aprova uma ocorrência pendente, mudando o status para 'Registrada'."""
    ocorrencia = db.get_or_404(Ocorrencia, ocorrencia_id)
    
    if ocorrencia.status != "Pendente":
        flash("Apenas ocorrências pendentes podem ser aprovadas.", "warning")
        return redirect(url_for("ocorrencia.detalhes_ocorrencia", ocorrencia_id=ocorrencia.id))
    
    try:
        ocorrencia.status = "Registrada"
        db.session.commit()
        flash("Ocorrência aprovada com sucesso!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Erro ao aprovar ocorrência: {e}", "danger")
    
    return redirect(url_for("ocorrencia.detalhes_ocorrencia", ocorrencia_id=ocorrencia.id))


@ocorrencia_bp.route("/rejeitar/<int:ocorrencia_id>", methods=["POST"])
@login_required
def rejeitar_ocorrencia(ocorrencia_id):
    """Rejeita uma ocorrência pendente, mudando o status para 'Rejeitada'."""
    ocorrencia = db.get_or_404(Ocorrencia, ocorrencia_id)
    
    if ocorrencia.status != "Pendente":
        flash("Apenas ocorrências pendentes podem ser rejeitadas.", "warning")
        return redirect(url_for("ocorrencia.detalhes_ocorrencia", ocorrencia_id=ocorrencia.id))
    
    try:
        ocorrencia.status = "Rejeitada"
        db.session.commit()
        flash("Ocorrência rejeitada.", "info")
    except Exception as e:
        db.session.rollback()
        flash(f"Erro ao rejeitar ocorrência: {e}", "danger")
    
    return redirect(url_for("ocorrencia.listar_ocorrencias"))


@ocorrencia_bp.route("/analisar-relatorio", methods=["POST", "OPTIONS"])
@cross_origin()
def analisar_relatorio():
    from app.services.ocorrencia_parser import OcorrenciaParser

    if request.method == "OPTIONS":
        return '', 200
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"erro": "Dados JSON não fornecidos"}), 400
        
        # Aceita tanto 'relatorio_bruto' quanto 'texto_relatorio' para compatibilidade
        texto = data.get("relatorio_bruto") or data.get("texto_relatorio", "")
        formatar_para_email = data.get("formatar_para_email", False)
        
        if not texto:
            return jsonify({"erro": "Texto do relatório está vazio"}), 400

        texto_limpo = texto.replace("\xa0", " ").strip()
        print(f"Texto recebido para análise: {texto_limpo[:200]}...")
        
        # Processamento e correção do texto
        relatorio_processado = OcorrenciaParser.processar_e_corrigir_texto(texto_limpo)
        print(f"Texto processado: {relatorio_processado[:200]}...")
        
        # Extrair dados estruturados do relatório
        dados_extraidos = OcorrenciaParser.extrair_dados_relatorio(relatorio_processado)
        print(f"Dados extraídos: {dados_extraidos}")
        
        # Preparar resposta
        resposta = {
            "relatorio_processado": relatorio_processado,
            "sucesso": True,
            "dados": dados_extraidos  # ✅ Adiciona os dados estruturados
        }
        
        # Verificação de segurança para garantir que os dados sejam válidos
        if not isinstance(dados_extraidos, dict):
            print(f"AVISO: dados_extraidos não é um dicionário: {type(dados_extraidos)}")
            dados_extraidos = {}
            resposta["dados"] = dados_extraidos
        
        # Se solicitado, gerar versão para email
        if formatar_para_email:
            relatorio_email = OcorrenciaParser.formatar_para_email_profissional(relatorio_processado)
            resposta["relatorio_email"] = relatorio_email
        
        return jsonify(resposta), 200
        
    except Exception as e:
        return jsonify({"erro": f"Erro interno: {str(e)}"}), 500


@ocorrencia_bp.route("/gerar-email-patrimonial/<int:ocorrencia_id>", methods=["POST"])
@login_required
def gerar_email_patrimonial(ocorrencia_id):
    """
    Rota para gerar a versão em formato de e-mail de um relatório patrimonial já salvo.
    Usa o serviço de IA dedicado com o prompt exclusivo.
    """
    ocorrencia = db.get_or_404(Ocorrencia, ocorrencia_id)

    if not ocorrencia.relatorio_final:
        return jsonify({"erro": "A ocorrência não possui um relatório final para formatar."}), 400

    try:
        service = EmailPatrimonialFormatService()
        email_formatado = service.formatar_email_patrimonial(ocorrencia.relatorio_final)

        return jsonify({
            "sucesso": True,
            "email_formatado": email_formatado
        }), 200

    except Exception as e:
        logger.error(f"Erro ao gerar e-mail patrimonial para ocorrência {ocorrencia.id}: {e}", exc_info=True)
        return jsonify({"erro": f"Erro ao gerar e-mail: {str(e)}"}), 500


@ocorrencia_bp.route("/gerar-email-consolidado", methods=["POST"])
@login_required
def gerar_email_consolidado():
    """
    Rota para compilar múltiplos relatórios patrimoniais em um único e-mail consolidado.
    """
    dados = request.get_json()
    if not dados or "ids" not in dados:
        return jsonify({"erro": "IDs não fornecidos."}), 400

    ocorrencia_ids = dados["ids"]
    if not ocorrencia_ids:
        return jsonify({"erro": "Nenhum ID de ocorrência fornecido."}), 400

    # Busca as ocorrências
    ocorrencias = Ocorrencia.query.filter(Ocorrencia.id.in_(ocorrencia_ids)).order_by(Ocorrencia.data_hora_ocorrencia.asc()).all()

    if not ocorrencias:
        return jsonify({"erro": "Nenhuma ocorrência encontrada para os IDs fornecidos."}), 404

    # Verifica se todas pertencem ao mesmo condomínio (segurança adicional)
    primeiro_condominio = ocorrencias[0].condominio.nome if ocorrencias[0].condominio else None
    for occ in ocorrencias:
        if (occ.condominio.nome if occ.condominio else None) != primeiro_condominio:
            return jsonify({"erro": "Não é possível consolidar relatórios de condomínios diferentes."}), 400

    # Prepara o texto unificado enviando os blocos
    blocos_texto = []
    for occ in ocorrencias:
        horario = occ.data_hora_ocorrencia.strftime('%H:%M') if occ.data_hora_ocorrencia else "S/H"
        relatorio = occ.relatorio_final or "Sem relatório final."
        blocos_texto.append(f"--- HORÁRIO DA OCORRÊNCIA: {horario} ---\n{relatorio}\n")

    texto_unificado = "\n".join(blocos_texto)

    try:
        service = EmailPatrimonialFormatService()
        email_formatado = service.formatar_email_consolidado(texto_unificado)

        return jsonify({
            "sucesso": True,
            "email_formatado": email_formatado
        }), 200

    except Exception as e:
        logger.error(f"Erro ao gerar e-mail consolidado: {e}", exc_info=True)
        return jsonify({"erro": f"Erro interno ao gerar consolidação: {str(e)}"}), 500

@ocorrencia_bp.route("/exportar-docx-consolidado", methods=["POST"])
@login_required
def exportar_docx_consolidado():
    """
    Rota para compilar múltiplos relatórios patrimoniais em um único arquivo DOCX consolidado.
    """
    dados = request.get_json()
    if not dados or "ids" not in dados:
        return jsonify({"erro": "IDs não fornecidos."}), 400

    ocorrencia_ids = dados["ids"]
    if not ocorrencia_ids:
        return jsonify({"erro": "Nenhum ID de ocorrência fornecido."}), 400

    # Busca as ocorrências
    ocorrencias = Ocorrencia.query.filter(Ocorrencia.id.in_(ocorrencia_ids)).order_by(Ocorrencia.data_hora_ocorrencia.asc()).all()

    if not ocorrencias:
        return jsonify({"erro": "Nenhuma ocorrência encontrada para os IDs fornecidos."}), 404

    try:
        from docx import Document
        from docx.shared import Pt, Cm, Mm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Criar documento Word
        document = Document()
        
        # Configurações da página A4 e margens moderadas (2cm)
        for section in document.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Estilos globais
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(0) # Impede espaços gigantescos entre parágrafos
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Título principal
        titulo = document.add_heading('3. OCORRÊNCIAS REGISTRADAS', level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in titulo.runs:
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0, 0, 0) # Força cor preta no título
        
        for idx, o in enumerate(ocorrencias, start=1):
            # Subtítulo da ocorrência
            subtitulo = document.add_heading(f'3.{idx} Ocorrência', level=2)
            for run in subtitulo.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0, 0, 0) # Força cor preta no subtítulo
            
            if o.relatorio_final:
                # O relatorio_final já vem formatado pela IA com todos os campos.
                linhas = o.relatorio_final.split('\n')
                in_relato = False
                
                current_p = None
                current_alignment = None
                
                for linha in linhas:
                    linha_texto = linha.rstrip('\r\n')
                    linha_limpa = linha_texto.strip()
                    
                    if not linha_limpa:
                        # Linha em branco estrutural original
                        document.add_paragraph()
                        current_p = None
                        continue
                        
                    # Controle de contexto para aplicar alinhamento justificado apenas no corpo do relato
                    if linha_limpa.startswith('Relato:'):
                        in_relato = True
                    elif linha_limpa.startswith('Ações Realizadas:') or \
                         linha_limpa.startswith('Acionamentos:') or \
                         linha_limpa.startswith('Envolvidos/Testemunhas:') or \
                         linha_limpa.startswith('Veículo') or \
                         linha_limpa.startswith('Responsável pelo registro:'):
                        in_relato = False

                    line_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if (in_relato and not linha_limpa.startswith('Relato:')) else WD_ALIGN_PARAGRAPH.LEFT

                    # Sempre cria novo parágrafo se for texto justificado (evita esticar linhas curtas).
                    # Se for alinhado à esquerda, agrupa na mesma caixa usando Shift+Enter (\n) para ficar bem colado.
                    force_new_p = (line_alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)

                    if current_p is None or current_alignment != line_alignment or force_new_p:
                        current_p = document.add_paragraph()
                        current_p.paragraph_format.space_after = Pt(0)
                        current_p.paragraph_format.line_spacing = 1.15
                        current_p.alignment = line_alignment
                        current_alignment = line_alignment
                    else:
                        # Adiciona quebra de linha manual dentro do mesmo parágrafo (fica perfeitamente agrupado)
                        current_p.add_run('\n')
                        
                    # Procura pelo primeiro dois-pontos para identificar o rótulo
                    parts = linha_texto.split(':', 1)
                    if len(parts) == 2 and len(parts[0]) < 40 and not linha_limpa.startswith('-'):
                        current_p.add_run(parts[0] + ':').bold = True
                        current_p.add_run(parts[1])
                    else:
                        current_p.add_run(linha_texto)
            else:
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run('Relatório final não disponível.').italic = True
            
            # Adicionar linha de separação entre ocorrências se não for a última
            if idx < len(ocorrencias):
                document.add_paragraph() # Espaço antes do separador
                
                p_sep = document.add_paragraph()
                p_sep.paragraph_format.space_after = Pt(0)
                p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Uma linha contínua ocupando a largura da área útil
                p_sep.add_run('_' * 78)
                
                document.add_paragraph() # Espaço depois do separador
        
        # Salvar em BytesIO
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        
        hoje_str = datetime.now().strftime('%d%m%Y')
        filename = f"Relatorio_consolidado_{hoje_str}.docx"
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"Erro ao exportar DOCX consolidado: {e}", exc_info=True)
        return jsonify({"erro": f"Erro interno ao gerar DOCX: {str(e)}"}), 500

